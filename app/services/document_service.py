import os
import csv
import logging
from typing import Tuple, Optional
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from defusedxml.ElementTree import ParseError

from app.models.document import DocumentContent

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
MAX_PDF_PAGES = 50
MAX_EXCEL_SHEETS = 20
MAX_EXCEL_ROWS = 3000
MAX_EXCEL_COLS = 100
MAX_EXTRACTED_TEXT_LENGTH = 60000


class DocumentService:
    def validate_document(self, file_path: str) -> Tuple[bool, str]:
        path = Path(file_path)
        if not path.exists():
            return False, "Файл не найден."
        if not path.is_file():
            return False, "Путь не является файлом."
        if path.stat().st_size == 0:
            return False, "Файл пустой."
        if path.stat().st_size > MAX_FILE_SIZE:
            return False, "Файл слишком большой (максимум 20 МБ)."
        ext = path.suffix.lower()
        if ext not in [".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".doc", ".xls"]:
            return False, f"Неподдерживаемый формат файла: {ext}"
        if ext in [".doc", ".xls"]:
            return False, "Старые форматы DOC и XLS пока не поддерживаются. Сохраните файл как DOCX или XLSX."
        try:
            with open(file_path, "rb") as f:
                f.read(4)  # Проверка возможности чтения
        except Exception:
            return False, "Нет прав на чтение файла или файл повреждён."
        return True, "Файл валиден."

    def extract_document(self, file_path: str) -> DocumentContent:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self.extract_pdf(file_path)
        elif ext == ".docx":
            return self.extract_docx(file_path)
        elif ext in [".xlsx", ".xlsm"]:
            return self.extract_excel(file_path)
        elif ext == ".csv":
            return self.extract_csv(file_path)
        elif ext == ".txt":
            return self.extract_txt(file_path)
        else:
            raise ValueError("Неподдерживаемый формат файла.")

    def extract_pdf(self, file_path: str) -> DocumentContent:
        warnings = []
        extracted_text = []
        page_count = 0
        try:
            reader = PdfReader(file_path)
            if reader.is_encrypted:
                warnings.append("PDF защищён паролем. Сначала снимите защиту или сохраните незашифрованную копию.")
                return DocumentContent(
                    file_path=file_path,
                    file_name=os.path.basename(file_path),
                    file_type="pdf",
                    extracted_text="",
                    page_count=0,
                    warnings=warnings,
                )
            page_count = min(len(reader.pages), MAX_PDF_PAGES)
            text_pages = 0
            for i in range(page_count):
                page = reader.pages[i]
                text = page.extract_text() or ""
                if not text.strip():
                    warnings.append(f"Страница {i+1} не содержит извлекаемого текста.")
                else:
                    text_pages += 1
                extracted_text.append(f"--- Страница {i+1} ---\n{text.strip()}")
            if text_pages / page_count < 0.3:
                warnings.append("В документе мало извлекаемого текста. Возможно, это сканированный PDF. Попробуйте загрузить нужную страницу как изображение.")
            full_text = "\n".join(extracted_text)
            if len(full_text) > MAX_EXTRACTED_TEXT_LENGTH:
                full_text = full_text[:MAX_EXTRACTED_TEXT_LENGTH]
                warnings.append("Извлечённый текст был обрезан из-за ограничения длины.")
            return DocumentContent(
                file_path=file_path,
                file_name=os.path.basename(file_path),
                file_type="pdf",
                extracted_text=full_text,
                page_count=page_count,
                warnings=warnings,
            )
        except Exception as e:
            logger.error(f"Ошибка при извлечении PDF: {e}")
            raise

    def extract_docx(self, file_path: str) -> DocumentContent:
        warnings = []
        extracted_text = []
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                extracted_text.append(para.text)
            for i, table in enumerate(doc.tables):
                extracted_text.append(f"[Таблица {i+1}]")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    extracted_text.append(row_text)
            full_text = "\n".join(extracted_text)
            if len(full_text) > MAX_EXTRACTED_TEXT_LENGTH:
                full_text = full_text[:MAX_EXTRACTED_TEXT_LENGTH]
                warnings.append("Извлечённый текст был обрезан из-за ограничения длины.")
            # Проверка на наличие изображений (не извлекаем текст из них)
            if any(shape.shape_type == 13 for shape in doc.inline_shapes):
                warnings.append("Изображения внутри DOCX не были распознаны как текст.")
            return DocumentContent(
                file_path=file_path,
                file_name=os.path.basename(file_path),
                file_type="docx",
                extracted_text=full_text,
                warnings=warnings,
            )
        except Exception as e:
            logger.error(f"Ошибка при извлечении DOCX: {e}")
            raise

    def extract_excel(self, file_path: str) -> DocumentContent:
        warnings = []
        extracted_text = []
        sheet_names = []
        try:
            wb = load_workbook(filename=file_path, read_only=True, data_only=True, keep_links=False)
            sheet_names = wb.sheetnames[:MAX_EXCEL_SHEETS]
            for sheet_name in sheet_names:
                ws = wb[sheet_name]
                extracted_text.append(f"--- Лист: {sheet_name} ---")
                row_count = 0
                for row in ws.iter_rows(max_row=MAX_EXCEL_ROWS, max_col=MAX_EXCEL_COLS, values_only=True):
                    if all(cell is None or str(cell).strip() == "" for cell in row):
                        continue
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                extracted_text.append(row_text)
                row_count += 1
                if row_count >= MAX_EXCEL_ROWS:
                    warnings.append("Данные были обрезаны из-за ограничения количества строк.")
                    break
            wb.close()
            full_text = "\n".join(extracted_text)
            if len(full_text) > MAX_EXTRACTED_TEXT_LENGTH:
                full_text = full_text[:MAX_EXTRACTED_TEXT_LENGTH]
                warnings.append("Извлечённый текст был обрезан из-за ограничения длины.")
            return DocumentContent(
                file_path=file_path,
                file_name=os.path.basename(file_path),
                file_type="excel",
                extracted_text=full_text,
                sheet_names=sheet_names,
                warnings=warnings,
            )
        except InvalidFileException:
            raise ValueError("Файл Excel повреждён или имеет неподдерживаемый формат.")
        except Exception as e:
            logger.error(f"Ошибка при извлечении Excel: {e}")
            raise

    def extract_csv(self, file_path: str) -> DocumentContent:
        warnings = []
        extracted_text = []
        try:
            with open(file_path, "rb") as f:
                raw = f.read(4096)
            encodings = ["utf-8", "utf-8-sig", "windows-1251", "latin-1"]
            encoding = None
            for enc in encodings:
                try:
                    raw.decode(enc)
                    encoding = enc
                    break
                except UnicodeDecodeError:
                    continue
            if encoding is None:
                raise ValueError("Невозможно определить кодировку CSV файла.")
            with open(file_path, encoding=encoding, newline='') as f:
                sample = f.read(2048)
                f.seek(0)
                sniffer = csv.Sniffer()
                dialect = sniffer.sniff(sample, delimiters=[',',';','\\t','|'])
                reader = csv.reader(f, dialect)
                row_count = 0
                for row in reader:
                    if row_count >= 3000:
                        warnings.append("Данные были обрезаны из-за ограничения количества строк.")
                        break
                    if len(row) > 100:
                        warnings.append("Данные были обрезаны из-за ограничения количества столбцов.")
                        row = row[:100]
                    extracted_text.append(" | ".join(row))
                    row_count += 1
            full_text = "\n".join(extracted_text)
            if len(full_text) > MAX_EXTRACTED_TEXT_LENGTH:
                full_text = full_text[:MAX_EXTRACTED_TEXT_LENGTH]
                warnings.append("Извлечённый текст был обрезан из-за ограничения длины.")
            return DocumentContent(
                file_path=file_path,
                file_name=os.path.basename(file_path),
                file_type="csv",
                extracted_text=full_text,
                warnings=warnings,
            )
        except Exception as e:
            logger.error(f"Ошибка при извлечении CSV: {e}")
            raise

    def extract_txt(self, file_path: str) -> DocumentContent:
        warnings = []
        try:
            encodings = ["utf-8", "utf-8-sig", "windows-1251"]
            encoding = None
            for enc in encodings:
                try:
                    with open(file_path, encoding=enc) as f:
                        text = f.read()
                    encoding = enc
                    break
                except UnicodeDecodeError:
                    continue
            if encoding is None:
                raise ValueError("Невозможно определить кодировку TXT файла.")
            text = self._clean_text(text)
            if len(text) > MAX_EXTRACTED_TEXT_LENGTH:
                text = text[:MAX_EXTRACTED_TEXT_LENGTH]
                warnings.append("Извлечённый текст был обрезан из-за ограничения длины.")
            return DocumentContent(
                file_path=file_path,
                file_name=os.path.basename(file_path),
                file_type="txt",
                extracted_text=text,
                warnings=warnings,
            )
        except Exception as e:
            logger.error(f"Ошибка при извлечении TXT: {e}")
            raise

    def _clean_text(self, text: str) -> str:
        # Удалить нулевые байты
        text = text.replace('\\x00', '')
        # Нормализовать переносы строк
        text = text.replace('\\r\\n', '\\n').replace('\\r', '\\n')
        # Убрать чрезмерные повторения пустых строк (не более 2 подряд)
        lines = text.split('\\n')
        cleaned_lines = []
        empty_count = 0
        for line in lines:
            if line.strip() == '':
                empty_count += 1
                if empty_count <= 2:
                    cleaned_lines.append(line)
            else:
                empty_count = 0
                cleaned_lines.append(line)
        return '\\n'.join(cleaned_lines)

    def format_for_ai(self, document: DocumentContent) -> str:
        header = f"К задаче приложен документ.\\nТип документа: {document.file_type.upper()}\\nИмя файла: {document.file_name}\\n"
        content = document.extracted_text
        footer = "\\n===== КОНЕЦ СОДЕРЖИМОГО ДОКУМЕНТА ====="
        return f"{header}\\n===== НАЧАЛО СОДЕРЖИМОГО ДОКУМЕНТА =====\\n{content}\\n{footer}\\n"
